import os
import shutil
import uuid
import logging
import json
import fitz  # PyMuPDF
import pandas as pd
import numpy as np
from datetime import datetime
from dotenv import load_dotenv
from keybert import KeyBERT
from openai import OpenAI
from langdetect import detect
from googletrans import Translator
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from sklearn.metrics.pairwise import cosine_similarity

# Setup logging
logging.basicConfig(level=logging.INFO)
load_dotenv()

# Initialize globals
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
embedder = OpenAIEmbeddings(model="text-embedding-3-small")
kw_model = KeyBERT()
translator = Translator()

import spacy
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    # Download if not present (though usually better to do in docker/setup)
    from spacy.cli import download
    download("en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

USER_DIR = "user_memory"
os.makedirs(USER_DIR, exist_ok=True)

def detect_language(text):
    try:
        lang_code = detect(text)
        flags = {
            "en": "🇬🇧", "es": "🇪🇸", "fr": "🇫🇷", "de": "🇩🇪",
            "it": "🇮🇹", "pt": "🇵🇹", "hi": "🇮🇳", "unknown": "🌐"
        }
        return flags.get(lang_code, "🌐")
    except Exception:
        return "🌐"

def translate_to_english(text):
    try:
        lang_code = detect(text)
        if lang_code != "en" and lang_code != "unknown":
            try:
                translated = translator.translate(text, src=lang_code, dest="en")
                return translated.text
            except Exception:
                # fallback to OpenAI translate if googletrans fails
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Translate the following text to English precisely."},
                        {"role": "user", "content": text}
                    ],
                    temperature=0
                )
                return resp.choices[0].message.content.strip()
        return text
    except Exception:
        return text

def get_user_vector_path(session_id):
    run_id = str(uuid.uuid4())[:8]
    return os.path.join(USER_DIR, f"faiss_{session_id}_{run_id}")

def extract_text_from_pdf(file_bytes):
    pdf = fitz.open(stream=file_bytes, filetype="pdf")
    records = []
    for page_num, page in enumerate(pdf, start=1):
        lines = page.get_text("text").split("\n")
        for line_num, line in enumerate(lines, start=1):
            if line.strip():
                records.append({"page": page_num, "line": line_num, "text": line.strip()})
    pdf.close()
    return records

def extract_text_from_docx(file_bytes):
    from docx import Document as DocxDocument
    import io
    doc = DocxDocument(io.BytesIO(file_bytes))
    return [{"page": 1, "line": i + 1, "text": p.text.strip()} for i, p in enumerate(doc.paragraphs) if p.text.strip()]

def extract_text_from_excel(file_bytes):
    import io
    df = pd.read_excel(io.BytesIO(file_bytes)).dropna(how="all")
    df.columns = [str(c).strip() for c in df.columns]
    lines = [" | ".join(map(str, row)) for _, row in df.iterrows()]
    return [{"page": 1, "line": i + 1, "text": t} for i, t in enumerate(lines)]

def extract_text_from_txt(file_bytes):
    text = file_bytes.decode("utf-8")
    return [{"page": 1, "line": i+1, "text": l.strip()} for i, l in enumerate(text.split("\n")) if l.strip()]

def chunk_text(records):
    splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    docs = []
    for rec in records:
        for idx, chunk in enumerate(splitter.split_text(rec["text"])):
            docs.append(Document(page_content=chunk, metadata={"page": rec["page"], "line": rec["line"], "chunk_id": idx}))
    return docs

def build_vector_store(docs, path):
    if os.path.exists(path):
        shutil.rmtree(path)
    vs = FAISS.from_documents(docs, embedder)
    vs.save_local(path)
    return vs

def generate_dynamic_keywords(obligations):
    keywords_dict = {}
    for ob in obligations:
        try:
            keyphrases = kw_model.extract_keywords(ob, keyphrase_ngram_range=(1, 3), stop_words="english", top_n=5)
            doc = nlp(ob)
            noun_phrases = [chunk.text.lower() for chunk in doc.noun_chunks]
            cleaned = list({kp[0].lower() for kp in keyphrases} | set(noun_phrases))
            keywords_dict[ob] = cleaned
        except Exception:
            keywords_dict[ob] = []
    return keywords_dict

def query_rag(vs, obligation, auto_keywords, top_k=6):
    retriever = vs.as_retriever(search_type="similarity", search_kwargs={"k": top_k})
    docs = retriever.get_relevant_documents(obligation)
    
    if not docs:
        return {
            "obligation": obligation, "is_present": "No", "reason": "No relevant clauses retrieved.",
            "similarity_score": 0.0, "keyword_hits": [], "confidence": 0.0,
            "page": None, "line": None, "supporting_clauses": []
        }
    
    ob_emb = embedder.embed_query(obligation)
    doc_embs = [embedder.embed_query(d.page_content) for d in docs]
    sims = cosine_similarity([ob_emb], doc_embs)[0]
    best_idx = int(np.argmax(sims))
    best_doc = docs[best_idx]
    best_score = float(sims[best_idx])
    
    obligation_keywords = auto_keywords.get(obligation, [])
    keyword_hits = [kw for kw in obligation_keywords if kw.lower() in best_doc.page_content.lower()]
    keyword_ratio = len(keyword_hits) / max(1, len(obligation_keywords))
    
    # Cosine Similarity Only for Confidence
    confidence = round(best_score * 100, 1)
    
    # Enhanced Prompt with Strengthened Guidelines
    prompt = f"""
You are a multilingual contract compliance analyst. Translate internally to English.
Analyze the obligation and the relevant clauses. Return valid JSON ONLY.

Task: Determine if the 'Obligation' is fully present and agreed to in the 'Relevant Clauses'.

Guidelines:
1. **Reasonable Efforts to Achieve Result**: If the Obligation requires a specific result (e.g., "remedy infringement", "fix the product"), and the Contract commits to "reasonable commercial efforts" to achieve that result, this is ACCEPTABLE. Return "Yes". The contract may also offer alternatives "in lieu of" the primary remedy (e.g., securing licenses instead of fixing) as long as these alternatives achieve the same end result (non-infringement, continued use).

2. **Strict Remedy Matching for Guarantees**: If the Obligation demands a GUARANTEE with specific remedies (e.g., "must fix or replace"), and the Contract adds a REFUND/REIMBURSE option that allows termination of use (walking away), this undermines the guarantee. Return "No".

3. **Refund as Escape Clause**: If the Obligation requires continued use/access (implied by "guarantee", "must secure rights", "must replace"), but the Contract allows the vendor to simply refund money and terminate, this is NON-COMPLIANCE. Return "No".

Return JSON:
{{
  "is_present": "Yes" or "No",
  "reason": "short 1-2 sentence rationale",
  "suggestion": "If 'No', provide a specific clause suggestion to add or modify to achieve compliance. If 'Yes', return null."
}}

Obligation:
{obligation}

Relevant Clauses:
{chr(10).join([d.page_content for d in docs])}
"""
    llm_status, llm_reason = "No", ""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a concise contract compliance expert who replies in JSON. You only output Yes or No."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=400
        )
        res_text = resp.choices[0].message.content.strip()
        if "```json" in res_text:
            res_text = res_text.split("```json")[1].split("```")[0]
        elif "```" in res_text:
            res_text = res_text.split("```")[1]
        parsed = json.loads(res_text)
        llm_status = parsed.get("is_present", "No").strip()
        
        # Normalize to Title Case (Yes/No)
        if llm_status.lower() == "yes":
            llm_status = "Yes"
        elif llm_status.lower() == "no":
            llm_status = "No"
        
        # Enforce strict Yes/No
        if llm_status not in ["Yes", "No"]:
            llm_status = "No"
            
        llm_reason = parsed.get("reason", "").strip()
        llm_suggestion = parsed.get("suggestion", None)
        
        # Fix suggestion logic: null/None for Yes, actual suggestion for No
        if llm_status == "Yes":
            llm_suggestion = None
        elif not llm_suggestion or llm_suggestion == "null":
            llm_suggestion = "Consider adding explicit language to address this obligation."
    except Exception:
        llm_reason = "Reason could not be parsed from model response."
        llm_suggestion = "Could not generate suggestion due to error."
    
    # Final Status based on LLM (Semantic Analysis)
    final_status = llm_status
    
    # Removed strict similarity threshold to allow LLM reasoning (semantic match) to prevail
    # even if cosine similarity is low due to vocabulary differences.

    supporting_clauses = [f"[Page {d.metadata.get('page')} Line {d.metadata.get('line')}] {d.page_content[:250].strip()}" for d in docs]
    
    return {
        "obligation": obligation, "is_present": final_status, "reason": llm_reason,
        "similarity_score": round(best_score, 3), "keyword_hits": keyword_hits,
        "confidence": confidence, "page": best_doc.metadata.get("page"), "line": best_doc.metadata.get("line"),
        "supporting_clauses": supporting_clauses,
        "suggestion": llm_suggestion
    }

def analyze_contract(obligations_file_bytes, obligations_filename, contract_file_bytes, contract_filename, session_id):
    # 1. Load Obligations
    import io
    if obligations_filename.endswith(".csv"):
        df_ob = pd.read_csv(io.BytesIO(obligations_file_bytes)).dropna(how="all")
    else:
        df_ob = pd.read_excel(io.BytesIO(obligations_file_bytes)).dropna(how="all")
    
    df_ob.columns = [str(c).strip() for c in df_ob.columns]
    
    # Translate obligations
    df_ob["Language"] = df_ob.iloc[:, 0].astype(str).apply(lambda x: detect_language(x))
    df_ob["Obligation_English"] = df_ob.iloc[:, 0].astype(str).apply(lambda x: translate_to_english(x).strip())
    df_ob = df_ob[df_ob["Obligation_English"].str.strip() != ""].reset_index(drop=True)
    obligations = df_ob["Obligation_English"].tolist()
    
    # 2. Extract Contract Text
    if contract_filename.endswith(".pdf"):
        records = extract_text_from_pdf(contract_file_bytes)
    elif contract_filename.endswith(".docx"):
        records = extract_text_from_docx(contract_file_bytes)
    elif contract_filename.endswith(".xlsx"):
        records = extract_text_from_excel(contract_file_bytes)
    else:
        records = extract_text_from_txt(contract_file_bytes)
        
    # Translate contract text
    for rec in records:
        rec["text"] = translate_to_english(rec["text"]).strip()
        
    # 3. Build Vector Store
    docs = chunk_text(records)
    vector_path = get_user_vector_path(session_id)
    vs = build_vector_store(docs, vector_path)
    
    # 4. Generate Keywords
    auto_keywords = generate_dynamic_keywords(obligations)
    
    # 5. Run Analysis
    results = []
    for ob in obligations:
        results.append(query_rag(vs, ob, auto_keywords))
        
    # Cleanup vector store
    # shutil.rmtree(vector_path) # Optional: keep it for caching or delete
    
    # 6. Prepare Full Text for Preview Fallback
    full_text = "\n\n".join([r["text"] for r in records])

    return results, full_text
